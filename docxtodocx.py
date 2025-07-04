import os
import logging
from telegram import Update
from telegram.ext import Application, CommandHandler, MessageHandler, filters, ContextTypes
import docx

# --- Configuration ---
# PASTE YOUR TELEGRAM BOT TOKEN HERE
BOT_TOKEN ="8077706019:AAE5qdQ6i4IyNkTAzxCiBv-45xFJSnCWD9o" 
TABLES_PER_FILE = 30
DOWNLOAD_DIR = "downloads"

# --- Setup Logging ---
logging.basicConfig(
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    level=logging.INFO
)
logger = logging.getLogger(__name__)

# --- Helper Function to Clone a Table ---
def clone_table(table, new_doc):
    """
    Clones a table from a source document to a new document.
    Note: This primarily copies text content and basic structure.
    Complex formatting like merged cells might not be perfectly preserved.
    """
    # Create a new table with the same number of columns
    new_table = new_doc.add_table(rows=0, cols=len(table.columns))
    new_table.style = table.style # Attempt to copy table style

    # Copy rows and cell content
    for row in table.rows:
        new_row_cells = new_table.add_row().cells
        for i, cell in enumerate(row.cells):
            new_row_cells[i].text = cell.text
            # To copy paragraph formatting (like bold, italic), you'd need a more complex function
            # This implementation focuses on preserving the text and structure.

# --- Bot Command Handlers ---

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Sends a welcome message when the /start command is issued."""
    await update.message.reply_html(
        "👋 **Welcome to the DOCX Table Converter Bot!**\n\n"
        "Please send me one or more `.docx` files. When you are done uploading, "
        "use the <code>/convert</code> command to process them."
    )

async def handle_document(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Handles receiving a document, saves it, and prompts the user."""
    message = update.message
    if not message.document:
        return

    if not message.document.file_name.endswith('.docx'):
        await message.reply_text("⚠️ Please send only `.docx` files.")
        return

    # Initialize user data storage if it doesn't exist
    user_id = message.from_user.id
    if 'files' not in context.user_data:
        context.user_data['files'] = []

    try:
        # Download the file
        file = await message.document.get_file()
        file_path = os.path.join(DOWNLOAD_DIR, f"{user_id}_{message.document.file_name}")
        await file.download_to_drive(file_path)
        
        # Store the path for later processing
        context.user_data['files'].append(file_path)
        logger.info(f"User {user_id} uploaded file: {file_path}")

        await message.reply_text(
            "✅ File received.\n\n"
            "You can now send another `.docx` file or use the <b>/convert</b> command to process all uploaded files.",
            parse_mode='HTML'
        )
    except Exception as e:
        logger.error(f"Error handling document for user {user_id}: {e}")
        await message.reply_text("An error occurred while receiving your file. Please try again.")


async def convert(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Processes the stored files, extracts tables, and sends back new files."""
    user_id = update.message.from_user.id
    
    if 'files' not in context.user_data or not context.user_data['files']:
        await update.message.reply_text("You haven't sent any `.docx` files yet. Please send at least one file before using /convert.")
        return

    await update.message.reply_text("🔄 Processing your files... This may take a moment.")

    all_tables = []
    input_files = context.user_data['files']

    # 1. Extract all tables from all uploaded documents
    try:
        for file_path in input_files:
            logger.info(f"Reading tables from {file_path}")
            doc = docx.Document(file_path)
            all_tables.extend(doc.tables)
    except Exception as e:
        logger.error(f"Error reading docx file for user {user_id}: {e}")
        await update.message.reply_text("❌ A critical error occurred while reading one of your files. Please ensure they are not corrupted.")
        # Clean up
        for file in input_files:
            if os.path.exists(file):
                os.remove(file)
        context.user_data['files'] = []
        return
        
    if not all_tables:
        await update.message.reply_text("ℹ️ No tables were found in the document(s) you sent.")
        return

    logger.info(f"Found a total of {len(all_tables)} tables for user {user_id}.")

    # 2. Create new documents with 30 tables each
    output_files = []
    file_counter = 1
    try:
        for i in range(0, len(all_tables), TABLES_PER_FILE):
            chunk = all_tables[i:i + TABLES_PER_FILE]
            
            new_doc = docx.Document()
            new_doc.add_heading(f"Converted Tables - Part {file_counter}", level=1)
            new_doc.add_paragraph(f"This document contains {len(chunk)} tables.")
            
            for table in chunk:
                clone_table(table, new_doc)
                new_doc.add_paragraph() # Add some space between tables

            output_filename = os.path.join(DOWNLOAD_DIR, f"{user_id}_output_part_{file_counter}.docx")
            new_doc.save(output_filename)
            output_files.append(output_filename)
            file_counter += 1
            
        # 3. Send the generated files back to the user
        await update.message.reply_text(f"✅ Conversion complete! Found {len(all_tables)} tables. Sending you {len(output_files)} new file(s)...")
        for output_file in output_files:
            with open(output_file, 'rb') as f:
                await context.bot.send_document(chat_id=update.effective_chat.id, document=f)

    except Exception as e:
        logger.error(f"Error during conversion process for user {user_id}: {e}")
        await update.message.reply_text("❌ An error occurred during the conversion process.")
    
    finally:
        # 4. Cleanup: Remove all downloaded and generated files
        logger.info(f"Cleaning up files for user {user_id}")
        all_temp_files = input_files + output_files
        for file_path in all_temp_files:
            if os.path.exists(file_path):
                os.remove(file_path)
        
        # Reset user data for the next batch
        context.user_data['files'] = []


# --- Main Bot Execution ---
def main() -> None:
    """Start the bot."""
    if BOT_TOKEN == "YOUR_TELEGRAM_BOT_TOKEN":
        print("!!! ERROR: Please replace 'YOUR_TELEGRAM_BOT_TOKEN' with your actual bot token. !!!")
        return

    # Create the download directory if it doesn't exist
    if not os.path.exists(DOWNLOAD_DIR):
        os.makedirs(DOWNLOAD_DIR)

    # Create the Application and pass it your bot's token.
    application = Application.builder().token(BOT_TOKEN).build()

    # Register handlers
    application.add_handler(CommandHandler("start", start))
    application.add_handler(CommandHandler("convert", convert))
    application.add_handler(MessageHandler(filters.Document.ALL, handle_document))

    # Start the Bot
    print("Bot is running...")
    application.run_polling()


if __name__ == '__main__':
    main()
